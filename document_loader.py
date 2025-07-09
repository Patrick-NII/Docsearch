import os
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
import PyPDF2
import pypdf
from docx import Document
import pandas as pd
from PIL import Image
import pytesseract
import easyocr
import io
import base64

from config import settings

logger = logging.getLogger(__name__)

class AdvancedDocumentLoader:
    """Chargeur de documents avancé avec support multi-format"""
    
    def __init__(self):
        self.supported_extensions = set(settings.SUPPORTED_FORMATS)
        self.ocr_reader = None
        self._init_ocr()
    
    def _init_ocr(self):
        """Initialise le système OCR"""
        if settings.ENABLE_OCR:
            try:
                # EasyOCR pour une meilleure reconnaissance multi-langue
                self.ocr_reader = easyocr.Reader(settings.OCR_LANGUAGES)
                logger.info("OCR EasyOCR initialisé")
            except Exception as e:
                logger.warning(f"Impossible d'initialiser EasyOCR: {e}")
                self.ocr_reader = None
    
    def load_documents(self, source_dir: str = None) -> List[Dict[str, Any]]:
        """
        Charge tous les documents supportés d'un répertoire
        
        Args:
            source_dir: Chemin vers le répertoire source
            
        Returns:
            Liste des documents chargés
        """
        source_path = Path(source_dir or settings.SOURCE_DIR)
        if not source_path.exists():
            logger.warning(f"Le répertoire {source_path} n'existe pas")
            return []
        
        documents = []
        
        for file_path in source_path.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                try:
                    doc = self.load_single_document(file_path)
                    if doc:
                        documents.append(doc)
                        logger.info(f"Document chargé: {file_path.name}")
                    
                except Exception as e:
                    logger.error(f"Erreur lors du chargement de {file_path}: {e}")
        
        return documents
    
    def load_single_document(self, file_path: Path) -> Optional[Dict[str, Any]]:
        """
        Charge un document unique selon son type
        
        Args:
            file_path: Chemin vers le fichier
            
        Returns:
            Dictionnaire contenant le texte et les métadonnées
        """
        extension = file_path.suffix.lower()
        
        try:
            if extension == '.pdf':
                return self.load_pdf(file_path)
            elif extension in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff']:
                return self.load_image(file_path)
            elif extension in ['.docx', '.doc']:
                return self.load_word(file_path)
            elif extension in ['.xlsx', '.xls']:
                return self.load_excel(file_path)
            elif extension == '.csv':
                return self.load_csv(file_path)
            elif extension == '.txt':
                return self.load_txt(file_path)
            else:
                logger.warning(f"Format non supporté: {extension}")
                return None
                
        except Exception as e:
            logger.error(f"Erreur lors du chargement de {file_path}: {e}")
            return None
    
    def load_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Charge un fichier PDF avec extraction de texte avancée"""
        try:
            text = ""
            metadata = {
                "source": str(file_path),
                "filename": file_path.name,
                "file_type": "pdf",
                "file_size": file_path.stat().st_size
            }
            
            # Essayer d'abord avec pypdf (plus récent)
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = pypdf.PdfReader(file)
                    metadata["total_pages"] = len(pdf_reader.pages)
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text.strip():
                            text += f"Page {page_num + 1}:\n{page_text}\n\n"
                        else:
                            # Si pas de texte, essayer OCR sur cette page
                            if self.ocr_reader:
                                # Convertir la page en image et faire OCR
                                pass  # TODO: Implémenter OCR sur pages PDF
                    
            except Exception as e:
                logger.warning(f"pypdf échoué, essai avec PyPDF2: {e}")
                # Fallback vers PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata["total_pages"] = len(pdf_reader.pages)
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        text += f"Page {page_num + 1}:\n{page_text}\n\n"
            
            return {
                "text": text.strip(),
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Erreur lors du chargement du PDF {file_path}: {e}")
            raise
    
    def load_image(self, file_path: Path) -> Dict[str, Any]:
        """Charge une image et extrait le texte via OCR"""
        try:
            text = ""
            metadata = {
                "source": str(file_path),
                "filename": file_path.name,
                "file_type": "image",
                "file_size": file_path.stat().st_size
            }
            
            # Ouvrir l'image
            with Image.open(file_path) as img:
                metadata["image_size"] = img.size
                metadata["image_mode"] = img.mode
                
                # Essayer OCR avec EasyOCR d'abord
                if self.ocr_reader:
                    try:
                        results = self.ocr_reader.readtext(str(file_path))
                        for (bbox, text_detected, confidence) in results:
                            if confidence > 0.5:  # Seuil de confiance
                                text += text_detected + " "
                        metadata["ocr_confidence"] = confidence if results else 0
                    except Exception as e:
                        logger.warning(f"EasyOCR échoué: {e}")
                
                # Fallback vers Tesseract si EasyOCR échoue
                if not text.strip():
                    try:
                        text = pytesseract.image_to_string(img, lang='fra+eng')
                        metadata["ocr_method"] = "tesseract"
                    except Exception as e:
                        logger.warning(f"Tesseract échoué: {e}")
                        text = f"[Image: {file_path.name}] - Texte non extrait"
                        metadata["ocr_method"] = "failed"
            
            return {
                "text": text.strip(),
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Erreur lors du chargement de l'image {file_path}: {e}")
            raise
    
    def load_word(self, file_path: Path) -> Dict[str, Any]:
        """Charge un document Word"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            metadata = {
                "source": str(file_path),
                "filename": file_path.name,
                "file_type": "word",
                "file_size": file_path.stat().st_size,
                "paragraphs": len(doc.paragraphs)
            }
            
            return {
                "text": text.strip(),
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Erreur lors du chargement du document Word {file_path}: {e}")
            raise
    
    def load_excel(self, file_path: Path) -> Dict[str, Any]:
        """Charge un fichier Excel"""
        try:
            df = pd.read_excel(file_path, sheet_name=None)
            text = ""
            
            for sheet_name, sheet_df in df.items():
                text += f"Feuille: {sheet_name}\n"
                text += sheet_df.to_string(index=False) + "\n\n"
            
            metadata = {
                "source": str(file_path),
                "filename": file_path.name,
                "file_type": "excel",
                "file_size": file_path.stat().st_size,
                "sheets": list(df.keys())
            }
            
            return {
                "text": text.strip(),
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Erreur lors du chargement du fichier Excel {file_path}: {e}")
            raise
    
    def load_csv(self, file_path: Path) -> Dict[str, Any]:
        """Charge un fichier CSV"""
        try:
            df = pd.read_csv(file_path)
            text = f"Fichier CSV: {file_path.name}\n"
            text += df.to_string(index=False)
            
            metadata = {
                "source": str(file_path),
                "filename": file_path.name,
                "file_type": "csv",
                "file_size": file_path.stat().st_size,
                "rows": len(df),
                "columns": list(df.columns)
            }
            
            return {
                "text": text.strip(),
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Erreur lors du chargement du fichier CSV {file_path}: {e}")
            raise
    
    def load_txt(self, file_path: Path) -> Dict[str, Any]:
        """Charge un fichier texte"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            metadata = {
                "source": str(file_path),
                "filename": file_path.name,
                "file_type": "txt",
                "file_size": file_path.stat().st_size
            }
            
            return {
                "text": text.strip(),
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Erreur lors du chargement du fichier texte {file_path}: {e}")
            raise
    
    def load_from_base64(self, base64_data: str, filename: str) -> Optional[Dict[str, Any]]:
        """Charge un document depuis des données base64 (pour upload API)"""
        try:
            # Décoder les données base64
            file_data = base64.b64decode(base64_data)
            
            # Créer un fichier temporaire en mémoire
            file_obj = io.BytesIO(file_data)
            
            # Déterminer le type de fichier
            extension = Path(filename).suffix.lower()
            
            if extension == '.pdf':
                # Créer un PDF temporaire
                temp_path = Path(f"/tmp/{filename}")
                with open(temp_path, 'wb') as f:
                    f.write(file_data)
                result = self.load_pdf(temp_path)
                temp_path.unlink()  # Supprimer le fichier temporaire
                return result
            elif extension in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff']:
                # Créer une image temporaire
                temp_path = Path(f"/tmp/{filename}")
                with open(temp_path, 'wb') as f:
                    f.write(file_data)
                result = self.load_image(temp_path)
                temp_path.unlink()
                return result
            else:
                logger.warning(f"Format non supporté pour upload: {extension}")
                return None
                
        except Exception as e:
            logger.error(f"Erreur lors du chargement depuis base64: {e}")
            return None 